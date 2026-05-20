"""
FMEA 자동화 파이프라인
Step 1: ARXML zip 파싱 → Supabase 프로젝트 + fmea_items 생성
Step 2: DBC 파일 → signal_range 채우기
Step 3: 기존 프로젝트 S/O/D 교차 참조 복사 (정규화 매칭)
Step 4: Claude AI → 미채워진 S/O/D 생성
"""

import os, re, json, time, shutil, tempfile, zipfile, concurrent.futures
import xml.etree.ElementTree as ET
from pathlib import Path
from collections import defaultdict
import requests, urllib3

urllib3.disable_warnings()

# ── 환경변수 ───────────────────────────────────────────────────────────────────
SB_URL = os.getenv("SUPABASE_URL",  "https://itzgdbeiyvodhfhmvrfw.supabase.co")
SB_KEY = os.getenv("SUPABASE_KEY",  "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6Iml0emdkYmVpeXZvZGhmaG12cmZ3Iiwicm9sZSI6ImFub24iLCJpYXQiOjE3NzY3NDE2MzcsImV4cCI6MjA5MjMxNzYzN30.iqzQr-3Lqf1O4UHFe9euTLyeIyBXreLPoSzUdtEaNP8")
AI_KEY = os.getenv("ANTHROPIC_API_KEY", "")

SB_H = {"apikey": SB_KEY, "Authorization": f"Bearer {SB_KEY}",
        "Content-Type": "application/json", "Prefer": "return=representation"}

NS = "http://autosar.org/schema/r4.0"
NP = {"ar": NS}

MODES_SR = ["MORE", "LESS", "CORRUPT", "EARLY", "LATE"]
MODES_CS = ["CORRUPT", "EARLY", "LATE"]

DETAIL_TMPL = {
    "MORE":    "{var}이(가) 정상 범위를 초과하는 값을 출력/수신",
    "LESS":    "{var}이(가) 정상 범위 미만의 값을 출력/수신",
    "CORRUPT": "{var}이(가) 논리적으로 잘못된 값을 출력/수신",
    "EARLY":   "{var}이(가) 예상 시점보다 일찍 업데이트/발생됨",
    "LATE":    "{var}이(가) 예상 시점보다 늦게 업데이트됨",
}

CAN_RX_COMPONENTS = {"CtAp_RxMainCAN", "CtAp_RxSubCAN"}
SKIP_PREFIX = ("CtDcm_", "CtDem_", "CtWdgM", "CtNvM", "CtSbc",
               "CtSafetyLib", "CtAp_Demo", "CtIoHwAb", "CtCdd_")

BASE_RANGES = {
    "uint8": "0 ~ 255", "uint8_t": "0 ~ 255",
    "uint16": "0 ~ 65535", "uint32": "0 ~ 4294967295",
    "sint8": "-128 ~ 127", "sint16": "-32768 ~ 32767",
    "sint32": "-2147483648 ~ 2147483647",
    "boolean": "0(False) ~ 1(True)", "float": "IEEE 754 float",
}

# S/O/D 소스 프로젝트 (S/O/D가 많은 순으로 배치 — 앞쪽 우선)
SRC_PROJECT_IDS = [
    "a43d7d4e-b104-4e90-ab55-04c7b31aa3e7",  # SX3 (같은 플랫폼 — 최우선)
    "0715f883-d3a1-4ddd-8a3b-d3071da9ed3e",  # JG1 SBW SW FMEA
    "32cc5148-96b5-42c5-a3dc-2b8fa3b9f93e",  # GN7_FL
    "89dc5818-2435-4d09-a1a9-36aea664d11d",  # LQ2
    "70f74c19-66b6-4b2e-a2b3-1ee04dd1b101",  # TK1
]

COPY_FIELDS = ["severity", "occurrence", "detection",
               "effect_system", "preventive_action", "safety_mechanism_text"]

# 변수명 정규화 관련
_DE_PREFIXES = ("de", "dg", "di", "dv", "dp", "dm")
_NOISE_PARTS = {"i", "p", "b", "u1", "u8", "u16", "u32", "ctap", "ctdcm", "raw"}


# ── 변수명 정규화 ──────────────────────────────────────────────────────────────
def normalize_varname(vn: str) -> str:
    """
    서로 다른 명명 체계를 통일:
      DeIdtSta              → idtsta
      u1_IgnOnStaChk        → ignonstack
      CtApSBWSigSet_I_u1_IdtSta → idtsta
      CtApPButtonSet_PButtonFault → pbuttonfault
    """
    low = vn.lower()

    # De/Dg/Di 접두어 제거 (ARXML 명명)
    for pfx in _DE_PREFIXES:
        if low.startswith(pfx) and len(low) > len(pfx) + 2:
            low = low[len(pfx):]
            break

    # u1_ / u8_ 같은 타입 접두어 제거
    low = re.sub(r'^u\d+_', '', low)

    # CtAp_Comp_I_u1_Signal 형식 처리 — 의미있는 마지막 부분 추출
    parts = low.split("_")
    meaningful = [p for p in parts if len(p) > 2 and p not in _NOISE_PARTS]
    if len(meaningful) >= 2:
        return meaningful[-1]  # 마지막 의미 있는 부분
    if meaningful:
        return meaningful[0]
    return parts[-1] if parts else low


def build_sod_lookup(rows: list[dict]) -> dict:
    """
    {norm_varname: {failure_mode: row}} 형태의 정규화 lookup 구축
    같은 키에 여러 행이 있으면 RPN 최대값 우선
    """
    lookup: dict[str, dict[str, dict]] = defaultdict(dict)
    for row in rows:
        vn = row["variable_name"]
        fm = row["failure_mode"]
        norm = normalize_varname(vn)
        existing = lookup[norm].get(fm)
        if existing is None or (row.get("rpn") or 0) > (existing.get("rpn") or 0):
            lookup[norm][fm] = row
    return dict(lookup)


# ── Supabase 헬퍼 ──────────────────────────────────────────────────────────────
def sb_get(table, params):
    all_rows, offset = [], 0
    while True:
        r = requests.get(f"{SB_URL}/rest/v1/{table}",
                         headers={**SB_H, "Range": f"{offset}-{offset+999}"},
                         params=params, verify=False)
        data = r.json()
        if not isinstance(data, list) or not data:
            break
        all_rows.extend(data)
        if len(data) < 1000:
            break
        offset += 1000
    return all_rows


def sb_post(table, data):
    r = requests.post(f"{SB_URL}/rest/v1/{table}", headers=SB_H,
                      json=data, verify=False)
    r.raise_for_status()
    return r.json()


def sb_patch(table, params, data):
    r = requests.patch(f"{SB_URL}/rest/v1/{table}",
                       headers={**SB_H, "Prefer": "return=minimal"},
                       params=params, json=data, verify=False)
    if r.status_code not in (200, 204):
        print(f"PATCH error {r.status_code}: {r.text[:300]}")
    return r.status_code in (200, 204)


# ── ZIP 압축 해제 + 경로 탐색 ──────────────────────────────────────────────────
def extract_zip(zip_bytes: bytes) -> Path:
    tmp = Path(tempfile.mkdtemp())
    with zipfile.ZipFile(__import__("io").BytesIO(zip_bytes)) as zf:
        zf.extractall(tmp)
    return tmp


def find_project_paths(base: Path):
    swcd_candidates = list(base.rglob("Swcd_App"))
    if not swcd_candidates:
        raise FileNotFoundError(
            "Swcd_App 디렉터리를 찾을 수 없습니다. "
            "ARXML 프로젝트 폴더를 zip으로 압축해서 올려주세요."
        )
    swcd_app_dir = swcd_candidates[0]

    comp_candidates = list(base.rglob("RootComposition.arxml"))
    if not comp_candidates:
        raise FileNotFoundError("RootComposition.arxml을 찾을 수 없습니다.")
    composition_file = comp_candidates[0]

    autosar_dt = list(base.rglob("AUTOSAR_DataTypes.arxml"))
    app_dt = list(swcd_app_dir.glob("*DataTypes*.arxml"))
    dt_files = app_dt + autosar_dt

    return swcd_app_dir, composition_file, dt_files


# ── ARXML 파싱 ────────────────────────────────────────────────────────────────
def _sname(el):
    s = el.find("ar:SHORT-NAME", NP)
    return s.text.strip() if s is not None and s.text else ""


def _ref_full(el, tag):
    r = el.find(f"ar:{tag}", NP)
    return r.text.strip() if r is not None and r.text else ""


def parse_interfaces(arxml_files):
    if_map, full_map = {}, {}
    for f in arxml_files:
        try:
            root = ET.parse(f).getroot()
        except ET.ParseError:
            continue

        def walk(el, path_prefix):
            for child in el:
                local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                name = _sname(child) if local not in ("AR-PACKAGES", "ELEMENTS") else ""
                cur_path = f"{path_prefix}/{name}" if name else path_prefix

                if local == "SENDER-RECEIVER-INTERFACE":
                    signals = [_sname(de) for de in child.findall(".//ar:VARIABLE-DATA-PROTOTYPE", NP) if _sname(de)]
                    if_map[name] = {"type": "SR", "signals": signals}
                    full_map[cur_path] = name
                elif local == "CLIENT-SERVER-INTERFACE":
                    ops = [_sname(op) for op in child.findall(".//ar:CLIENT-SERVER-OPERATION", NP) if _sname(op)]
                    if_map[name] = {"type": "CS", "signals": ops}
                    full_map[cur_path] = name
                else:
                    walk(child, cur_path)

        walk(root, "")
    return if_map, full_map


def parse_swc_ports(arxml_files):
    comp_ports = {}
    for f in arxml_files:
        try:
            root = ET.parse(f).getroot()
        except ET.ParseError:
            continue
        for swc in root.findall(".//ar:APPLICATION-SW-COMPONENT-TYPE", NP):
            comp = _sname(swc)
            if not comp:
                continue
            ports = {}
            for p in swc.findall(".//ar:P-PORT-PROTOTYPE", NP):
                pn = _sname(p)
                ref = _ref_full(p, "PROVIDED-INTERFACE-TREF")
                if pn and ref:
                    ports[pn] = {"kind": "P", "if_ref": ref.split("/")[-1]}
            for r in swc.findall(".//ar:R-PORT-PROTOTYPE", NP):
                rn = _sname(r)
                ref = _ref_full(r, "REQUIRED-INTERFACE-TREF")
                if rn and ref:
                    ports[rn] = {"kind": "R", "if_ref": ref.split("/")[-1]}
            comp_ports[comp] = ports
    return comp_ports


def parse_connectors(composition_file):
    connectors = []
    root = ET.parse(composition_file).getroot()
    for conn in root.findall(".//ar:ASSEMBLY-SW-CONNECTOR", NP):
        prov = conn.find("ar:PROVIDER-IREF", NP) or ET.Element("")
        req  = conn.find("ar:REQUESTER-IREF", NP) or ET.Element("")
        pc = _ref_full(prov, "CONTEXT-COMPONENT-REF").split("/")[-1]
        pp = _ref_full(prov, "TARGET-P-PORT-REF").split("/")[-1]
        rc = _ref_full(req,  "CONTEXT-COMPONENT-REF").split("/")[-1]
        if pc and pp and rc:
            connectors.append({"provider_comp": pc, "provider_port": pp, "requester_comp": rc})
    return connectors


def generate_items(comp_ports, if_map, connectors):
    receiver_map = defaultdict(list)
    for c in connectors:
        key = (c["provider_comp"], c["provider_port"])
        if c["requester_comp"] not in receiver_map[key]:
            receiver_map[key].append(c["requester_comp"])

    items, item_no, seen = [], 1, set()
    for comp, ports in sorted(comp_ports.items()):
        if any(comp.startswith(p) for p in SKIP_PREFIX):
            continue
        is_can_rx = comp in CAN_RX_COMPONENTS

        for port_name, port_info in ports.items():
            if port_info["kind"] != "P":
                continue
            iface = if_map.get(port_info["if_ref"])
            if not iface:
                continue

            receivers = receiver_map.get((comp, port_name), [])
            effect_module = ", ".join(receivers) if receivers else None
            category = "External" if is_can_rx else "Internal"
            modes = MODES_CS if iface["type"] == "CS" else MODES_SR

            for signal in iface["signals"]:
                for mode in modes:
                    key = (comp, signal, mode)
                    if key in seen:
                        continue
                    seen.add(key)
                    items.append({
                        "comp": comp, "item_no": str(item_no),
                        "category": category,
                        "variable_name": signal, "variable_type": "uint8",
                        "failure_mode": mode,
                        "failure_detail": DETAIL_TMPL[mode].format(var=signal),
                        "effect_module": effect_module,
                        "potential_cause": ("CAN 통신 오류 / 외부 ECU 신호 이상"
                                           if category == "External"
                                           else "SW 연산 오류 / 내부 상태 데이터 이상"),
                    })
                    item_no += 1
    return items


# ── DBC 파싱 ──────────────────────────────────────────────────────────────────
def parse_dbc_all(dbc_paths):
    signals, vals = {}, {}
    for dbc_path in dbc_paths:
        try:
            with open(dbc_path, encoding="utf-8", errors="replace") as f:
                lines = f.readlines()
        except Exception:
            continue
        for line in lines:
            line = line.strip()
            m = re.match(r'SG_\s+(\w+)\s*:\s*\d+\|(\d+)@\d+([+-])\s*'
                         r'\(([^,]+),([^)]+)\)\s*\[([^|]*)\|([^\]]*)\]', line)
            if m:
                name, bit_len = m.group(1), int(m.group(2))
                is_signed = m.group(3) == "-"
                factor, offset = float(m.group(4)), float(m.group(5))
                raw_min = -(2**(bit_len-1)) if is_signed else 0
                raw_max = 2**(bit_len-1)-1  if is_signed else 2**bit_len-1
                phy_min = factor*raw_min + offset
                phy_max = factor*raw_max + offset
                try:
                    mn_f, mx_f = float(m.group(6)), float(m.group(7))
                    if mn_f != 0.0 or mx_f != 0.0:
                        phy_min, phy_max = mn_f, mx_f
                except ValueError:
                    pass
                if name not in signals:
                    signals[name] = {"phy_min": phy_min, "phy_max": phy_max}
            m2 = re.match(r'VAL_\s+\d+\s+(\w+)\s+(.*?)\s*;', line)
            if m2:
                sname = m2.group(1)
                enum_map = {int(vm.group(1)): vm.group(2)
                            for vm in re.finditer(r'(\d+)\s+"([^"]*)"', m2.group(2))}
                if enum_map and sname not in vals:
                    vals[sname] = enum_map
    return signals, vals


def make_range_str(sig_info, enum_map):
    if enum_map:
        items = sorted(enum_map.items())[:8]
        labels = ", ".join(f"{v}={l}" for v, l in items)
        if len(enum_map) > 8:
            labels += " ..."
        return f"[0 ~ {max(enum_map)}] {labels}"
    mn, mx = sig_info["phy_min"], sig_info["phy_max"]
    fmt = lambda x: str(int(x)) if x == int(x) else f"{x:.4g}"
    return f"{fmt(mn)} ~ {fmt(mx)}"


def build_dbc_lookup(signals):
    lookup = {}
    for sname in signals:
        low = sname.lower()
        lookup[low] = sname
        parts = low.split("_")
        for i in range(1, len(parts)):
            suffix = "_".join(parts[i:])
            if suffix not in lookup:
                lookup[suffix] = sname
    return lookup


def match_dbc(variable_name, dbc_lookup, signals, vals):
    candidates = [variable_name]
    for prefix in ("De", "Dg", "Di", "Dv", "Dp", "Dm"):
        if variable_name.startswith(prefix) and len(variable_name) > len(prefix):
            candidates.append(variable_name[len(prefix):])
    for cand in candidates:
        low = cand.lower()
        if low in dbc_lookup:
            sname = dbc_lookup[low]
            return make_range_str(signals[sname], vals.get(sname, {}))
        for key, sname in dbc_lookup.items():
            if len(low) > 4 and (low in key or key.endswith(low)):
                return make_range_str(signals[sname], vals.get(sname, {}))
    return None


def parse_arxml_types(swcd_app_dir: Path, dt_files):
    var_type = {}
    port_if = swcd_app_dir / "App_PortInterface.arxml"
    if port_if.exists():
        for vdp in ET.parse(port_if).getroot().iter(f'{{{NS}}}VARIABLE-DATA-PROTOTYPE'):
            n = vdp.find(f'{{{NS}}}SHORT-NAME')
            t = vdp.find(f'{{{NS}}}TYPE-TREF')
            if n is not None and t is not None:
                var_type[n.text] = t.text.split('/')[-1]

    custom_base = {}
    for dtpath in dt_files:
        if not Path(dtpath).exists():
            continue
        for idt in ET.parse(dtpath).getroot().iter(f'{{{NS}}}IMPLEMENTATION-DATA-TYPE'):
            n = idt.find(f'{{{NS}}}SHORT-NAME')
            if n is None:
                continue
            tname = n.text
            base_ref = idt.find(f'.//{{{NS}}}BASE-TYPE-REF')
            if base_ref is not None:
                custom_base[tname] = base_ref.text.split('/')[-1]
            lo = idt.find(f'.//{{{NS}}}LOWER-LIMIT')
            hi = idt.find(f'.//{{{NS}}}UPPER-LIMIT')
            if lo is not None and hi is not None and tname not in custom_base:
                custom_base[tname] = f"{lo.text} ~ {hi.text}"

    def resolve(type_name):
        if type_name in BASE_RANGES:
            return BASE_RANGES[type_name]
        base = custom_base.get(type_name, "")
        if base in BASE_RANGES:
            return BASE_RANGES[base]
        if "uint8" in base.lower():
            return BASE_RANGES["uint8"]
        if "uint16" in base.lower():
            return BASE_RANGES["uint16"]
        return None

    return {vn: resolve(tn) for vn, tn in var_type.items() if resolve(tn)}


# ── 규칙 기반 S/O/D ──────────────────────────────────────────────────────────
# 신호명 키워드 → Severity (안전 영향도)
_SEV_KEYWORDS = [
    (10, ("steering", "brake", "accel")),
    (9,  ("lvrpos", "gearpos", "shiftpos", "parkpos", "shiftact", "sbw")),
    (8,  ("ign", "ecumodeect", "drvrdyst", "lvr", "park", "busoff", "canbus")),
    (7,  ("hall", "sensor", "flt", "fault", "err", "idtflt", "idtsta")),
    (6,  ("bat", "vbat", "ldo", "power", "cal", "calib", "can", "rx", "tx")),
    (5,  ("sta", "status", "chk", "check", "mon", "info")),
]

# failure_mode → Detection 기본값 (검출 난이도)
_DET_BASE = {
    "MORE":    3,   # 범위 초과 → 범위 체크로 검출 쉬움
    "LESS":    3,
    "CORRUPT": 5,   # 논리 오류 → 패리티/CRC 필요
    "EARLY":   6,   # 타이밍 → 검출 어려움
    "LATE":    6,
}

# failure_mode → Occurrence 기본값
_OCC_BASE = {
    "MORE": 3, "LESS": 3, "CORRUPT": 3,
    "EARLY": 2, "LATE": 3,
}

def rule_based_sod(variable_name: str, failure_mode: str, category: str) -> dict:
    """신호명·실패모드·카테고리 기반 S/O/D 규칙 산출"""
    vn_low = variable_name.lower()

    # Severity: 키워드 매칭
    sev = 5  # 기본값
    for s_val, keywords in _SEV_KEYWORDS:
        if any(kw in vn_low for kw in keywords):
            sev = s_val
            break
    # External 신호는 +1 (외부 ECU 의존성)
    if category == "External" and sev < 10:
        sev = min(sev + 1, 10)

    # Occurrence
    occ = _OCC_BASE.get(failure_mode, 3)
    if category == "External":
        occ = min(occ + 1, 10)  # CAN 수신 신호는 발생 가능성 조금 높음

    # Detection
    det = _DET_BASE.get(failure_mode, 5)

    rpn = sev * occ * det

    # effect_system: 간단한 설명
    effect_map = {
        "MORE":    f"{variable_name} 과도한 값으로 인한 시스템 오동작",
        "LESS":    f"{variable_name} 부족한 값으로 인한 기능 미작동",
        "CORRUPT": f"{variable_name} 잘못된 값으로 인한 제어 오류",
        "STUCK":   f"{variable_name} 고착으로 인한 기능 상실",
        "EARLY":   f"{variable_name} 조기 발생으로 인한 순서 오류",
        "LATE":    f"{variable_name} 지연으로 인한 타임아웃 또는 기능 손실",
        "ERRATIC": f"{variable_name} 불규칙 변동으로 인한 시스템 불안정",
    }
    preventive_map = {
        "MORE":    "입력 신호 범위 검사 (range check) 구현, 소프트웨어 안전 메커니즘 적용",
        "LESS":    "최소값 임계치 모니터링, 신호 유효성 검증 로직 구현",
        "CORRUPT": "E2E 프로텍션 또는 CRC 검사 적용, 신호 논리 유효성 검증",
        "STUCK":   "신호 변화량 모니터링 (debounce + stuck detection), 주기 검사",
        "EARLY":   "타이밍 윈도우 검사, 이벤트 순서 검증 로직 구현",
        "LATE":    "타임아웃 모니터링 구현, 통신 주기 감시 (alive counter)",
        "ERRATIC": "이동 평균 필터 또는 디바운스 로직, 신호 안정성 검증",
    }

    return {
        "severity":          sev,
        "occurrence":        occ,
        "detection":         det,
        "rpn":               rpn,
        "effect_system":     effect_map.get(failure_mode, f"{variable_name} 신호 이상으로 인한 기능 영향"),
        "preventive_action": preventive_map.get(failure_mode, "신호 유효성 검증 및 안전 메커니즘 구현"),
        "ai_generated":      False,
    }


# ── Claude AI S/O/D (크레딧 있을 때만 동작) ──────────────────────────────────
def ai_analyze_batch(batch: list[dict], vehicle_model: str) -> dict:
    if not AI_KEY:
        return {}
    import anthropic
    client = anthropic.Anthropic(api_key=AI_KEY)

    items_text = "\n".join(
        f"{i+1}. comp={s['comp']}, var={s['variable_name']}, "
        f"mode={s['failure_mode']}, range={s.get('signal_range','?')}"
        for i, s in enumerate(batch)
    )
    prompt = (
        f"당신은 {vehicle_model} ECU SW FMEA 전문가입니다 (AIAG/VDA 기준).\n"
        f"아래 {len(batch)}개 항목의 S(Severity), O(Occurrence), D(Detection)(각 1~10)과\n"
        "effect_system, preventive_action을 한국어로 작성하세요.\n\n"
        f"{items_text}\n\n"
        "JSON 배열로만 응답 (마크다운 없이):\n"
        '[{"idx":1,"severity":7,"occurrence":3,"detection":4,'
        '"effect_system":"시스템 영향","preventive_action":"예방조치"},...]'
    )
    try:
        resp = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=4096,
            messages=[{"role": "user", "content": prompt}]
        )
        text = resp.content[0].text.strip()
        m = re.search(r'\[.*\]', text, re.DOTALL)
        if m:
            results = json.loads(m.group())
            out = {}
            for res in results:
                idx = res.get("idx", 0) - 1
                if 0 <= idx < len(batch):
                    key = (batch[idx]["comp"], batch[idx]["variable_name"])
                    out[key] = res
            return out
    except Exception as e:
        print(f"AI 오류: {e}")
    return {}


# ── 파이프라인 메인 ───────────────────────────────────────────────────────────
def run_pipeline(job_id, jobs, project_name, vehicle_model,
                 arxml_zip_bytes: bytes, dbc_data: list):
    job = jobs[job_id]
    tmp_dir = None

    def log(msg: str, progress: int | None = None):
        job["logs"].append(msg)
        if progress is not None:
            job["progress"] = progress
        print(f"[{job_id[:8]}] {msg}")

    try:
        # ── Step 0: 압축 해제 ──────────────────────────────────────────────────
        log("ARXML 압축 해제 중...", 3)
        tmp_dir = extract_zip(arxml_zip_bytes)
        swcd_app_dir, composition_file, dt_files = find_project_paths(tmp_dir)
        log(f"  구조 확인: {swcd_app_dir.name} / {composition_file.name}", 8)

        dbc_paths = []
        for fname, fdata in dbc_data:
            p = tmp_dir / fname
            p.write_bytes(fdata)
            dbc_paths.append(p)
        log(f"  DBC 파일: {len(dbc_paths)}개")

        # ── Step 1: ARXML 파싱 ────────────────────────────────────────────────
        log("ARXML 인터페이스 파싱...", 12)
        arxml_files = list(swcd_app_dir.glob("*.arxml"))
        if not arxml_files:
            raise FileNotFoundError("Swcd_App/*.arxml 파일이 없습니다.")

        if_map, _  = parse_interfaces(arxml_files)
        comp_ports = parse_swc_ports(arxml_files)
        connectors = parse_connectors(composition_file)
        log(f"  인터페이스: {len(if_map)}개 / 컴포넌트: {len(comp_ports)}개 / 커넥터: {len(connectors)}개", 18)

        items = generate_items(comp_ports, if_map, connectors)
        log(f"  FMEA 항목: {len(items)}개 생성", 24)

        # ── Supabase 프로젝트 생성 ─────────────────────────────────────────────
        log("Supabase 프로젝트 생성...", 26)
        existing = sb_get("projects", {"name": f"eq.{project_name}", "select": "id,name"})
        if existing:
            project_id = existing[0]["id"]
            requests.delete(f"{SB_URL}/rest/v1/fmea_items", headers=SB_H,
                            params={"project_id": f"eq.{project_id}"}, verify=False)
            requests.delete(f"{SB_URL}/rest/v1/sw_units", headers=SB_H,
                            params={"project_id": f"eq.{project_id}"}, verify=False)
            log(f"  기존 프로젝트 초기화: {project_id[:8]}...")
        else:
            r = sb_post("projects", {
                "name": project_name, "vehicle_model": vehicle_model,
                "description": f"SBW FMEA - {project_name} (자동 생성)"
            })
            project_id = (r[0] if isinstance(r, list) else r)["id"]
            log(f"  새 프로젝트: {project_id[:8]}...")

        # SW Units 생성
        unit_names = list({it["comp"] for it in items})
        unit_map = {}
        for name in unit_names:
            r = sb_post("sw_units", {"project_id": project_id, "name": name})
            uid = (r[0] if isinstance(r, list) else r)["id"]
            unit_map[name] = uid

        # fmea_items 일괄 삽입
        log(f"  fmea_items {len(items)}개 삽입 중...", 30)
        db_rows = [{
            "project_id":     project_id,
            "sw_unit_id":     unit_map.get(it["comp"]),
            "item_no":        it["item_no"],
            "category":       it["category"],
            "variable_name":  it["variable_name"],
            "variable_type":  it["variable_type"],
            "failure_mode":   it["failure_mode"],
            "failure_detail": it["failure_detail"],
            "effect_module":  it["effect_module"],
            "potential_cause": it["potential_cause"],
            "status":         "draft",
            "ai_generated":   False,
        } for it in items]

        inserted = 0
        for i in range(0, len(db_rows), 200):
            chunk = db_rows[i:i+200]
            r = requests.post(f"{SB_URL}/rest/v1/fmea_items", headers=SB_H,
                              json=chunk, verify=False)
            if r.status_code in (200, 201):
                inserted += len(chunk)
            log(f"  삽입 중: {min(i+200, len(db_rows))}/{len(db_rows)}")
        log(f"  삽입 완료: {inserted}개", 40)

        # ── Step 2: DBC → signal_range ────────────────────────────────────────
        log("DBC 신호 범위 매칭...", 42)
        if dbc_paths:
            signals, vals = parse_dbc_all(dbc_paths)
            log(f"  DBC 신호: {len(signals)}개 / VAL_: {len(vals)}개", 44)
            dbc_lookup   = build_dbc_lookup(signals)
            arxml_ranges = parse_arxml_types(swcd_app_dir, dt_files)

            var_range = {}
            for it in items:
                vn = it["variable_name"]
                if vn in var_range:
                    continue
                rng = match_dbc(vn, dbc_lookup, signals, vals)
                if not rng and vn in arxml_ranges:
                    rng = arxml_ranges[vn]
                if not rng:
                    rng = BASE_RANGES.get(it.get("variable_type", "").lower())
                var_range[vn] = rng

            to_update = {vn: rng for vn, rng in var_range.items() if rng}
            done_range = 0
            with concurrent.futures.ThreadPoolExecutor(max_workers=15) as ex:
                futures = {ex.submit(sb_patch, "fmea_items",
                                     {"project_id": f"eq.{project_id}",
                                      "variable_name": f"eq.{vn}"},
                                     {"signal_range": rng}): vn
                           for vn, rng in to_update.items()}
                for f in concurrent.futures.as_completed(futures):
                    if f.result():
                        done_range += 1
            log(f"  signal_range: {done_range}/{len(to_update)}개 변수 완료", 58)
        else:
            log("  DBC 없음, signal_range 스킵", 58)

        # ── Step 3: S/O/D 교차 참조 (정규화 매칭) ────────────────────────────
        log("기존 프로젝트 S/O/D 교차 참조 (정규화 매칭)...", 60)
        FIELD_STR = ",".join(["id", "variable_name", "failure_mode", "rpn"] + COPY_FIELDS)

        # 소스 프로젝트 로드 및 정규화 lookup 구축
        all_src_rows = []
        for pid in SRC_PROJECT_IDS:
            if pid == project_id:
                continue
            rows = sb_get("fmea_items", {
                "project_id": f"eq.{pid}",
                "severity":   "not.is.null",
                "select":     FIELD_STR,
            })
            all_src_rows.extend(rows)
            log(f"  소스 {pid[:8]}: {len(rows)}개")

        # 정규화 lookup 구축
        norm_lookup = build_sod_lookup(all_src_rows)
        log(f"  정규화 lookup: {len(norm_lookup)}개 고유 신호명 키", 65)

        # 타겟 항목 조회
        sx3_todo = sb_get("fmea_items", {
            "project_id": f"eq.{project_id}",
            "severity":   "is.null",
            "select":     "id,variable_name,failure_mode",
        })
        log(f"  S/O/D 미입력: {len(sx3_todo)}개")

        # 매칭
        crossref_updates = []
        for item in sx3_todo:
            vn   = item["variable_name"]
            fm   = item["failure_mode"]
            norm = normalize_varname(vn)

            src_row = None
            # 1. 정규화 정확 매칭
            if norm in norm_lookup and fm in norm_lookup[norm]:
                src_row = norm_lookup[norm][fm]
            # 2. 정규화 부분 매칭 (같은 failure_mode)
            if src_row is None:
                for key, fm_map in norm_lookup.items():
                    if fm in fm_map and len(key) > 3 and (key in norm or norm in key):
                        src_row = fm_map[fm]
                        break

            if src_row:
                patch = {f: src_row[f] for f in COPY_FIELDS if src_row.get(f) is not None}
                if patch:
                    crossref_updates.append((item["id"], patch))

        log(f"  매칭: {len(crossref_updates)}/{len(sx3_todo)}개 ({len(crossref_updates)/max(len(sx3_todo),1)*100:.0f}%)", 68)

        done_cr = 0
        with concurrent.futures.ThreadPoolExecutor(max_workers=15) as ex:
            futs = {ex.submit(sb_patch, "fmea_items", {"id": f"eq.{iid}"}, patch): iid
                    for iid, patch in crossref_updates}
            for f in concurrent.futures.as_completed(futs):
                if f.result():
                    done_cr += 1
        log(f"  교차 참조 완료: {done_cr}개 업데이트", 75)

        # ── Step 4: 규칙 기반 S/O/D → 나머지 미입력 항목 ────────────────────
        log("규칙 기반 S/O/D 생성...", 77)
        rule_todo = sb_get("fmea_items", {
            "project_id": f"eq.{project_id}",
            "severity":   "is.null",
            "select":     "id,variable_name,failure_mode,category",
        })
        log(f"  규칙 적용 대상: {len(rule_todo)}개 항목")

        done_rule = 0
        with concurrent.futures.ThreadPoolExecutor(max_workers=15) as ex:
            def apply_rule(row):
                sod = rule_based_sod(
                    row["variable_name"],
                    row["failure_mode"] or "MORE",
                    row["category"] or "Internal",
                )
                patch = {
                    "severity":          sod["severity"],
                    "occurrence":        sod["occurrence"],
                    "detection":         sod["detection"],
                    "effect_system":     sod["effect_system"],
                    "preventive_action": sod["preventive_action"],
                }
                return sb_patch("fmea_items", {"id": f"eq.{row['id']}"}, patch)

            futures = [ex.submit(apply_rule, row) for row in rule_todo]
            for f in concurrent.futures.as_completed(futures):
                if f.result():
                    done_rule += 1

        log(f"  규칙 기반 완료: {done_rule}개 항목 업데이트", 97)

        # ── 완료 ─────────────────────────────────────────────────────────────
        # 최종 통계
        final = sb_get("fmea_items", {
            "project_id": f"eq.{project_id}",
            "select": "signal_range,severity",
        })
        n_range = sum(1 for i in final if i.get("signal_range"))
        n_sod   = sum(1 for i in final if i.get("severity"))
        log(f"  최종: signal_range={n_range}/{len(final)}  S/O/D={n_sod}/{len(final)}", 99)

        job["project_id"] = project_id
        job["status"]     = "done"
        job["progress"]   = 100
        log(f"완료! {project_name}: signal_range={n_range}, S/O/D={n_sod}/{len(final)}")

    except Exception as e:
        import traceback
        job["status"] = "error"
        job["error"]  = str(e)
        log(f"오류: {e}")
        traceback.print_exc()
    finally:
        if tmp_dir and tmp_dir.exists():
            shutil.rmtree(tmp_dir, ignore_errors=True)


# ── 프로젝트 비교 ────────────────────────────────────────────────────────────
_DIFF_ORDER = {"b_missing": 0, "different": 1, "both_missing": 2, "same": 3,
               "a_only": 4, "b_only": 5}

def compare_fmea_projects(proj_a_id: str, proj_b_id: str) -> dict:
    """두 프로젝트 FMEA 항목을 정규화 매칭으로 비교"""
    FIELD_STR = "id,variable_name,failure_mode,severity,occurrence,detection,rpn,effect_system,preventive_action"

    rows_a = sb_get("fmea_items", {"project_id": f"eq.{proj_a_id}", "select": FIELD_STR})
    rows_b = sb_get("fmea_items", {"project_id": f"eq.{proj_b_id}", "select": FIELD_STR})

    def build_lookup(rows):
        lookup: dict[tuple, dict] = {}
        for row in rows:
            norm = normalize_varname(row["variable_name"])
            key = (norm, row["failure_mode"])
            if key not in lookup or (row.get("rpn") or 0) > (lookup[key].get("rpn") or 0):
                lookup[key] = row
        return lookup

    lookup_a = build_lookup(rows_a)
    lookup_b = build_lookup(rows_b)

    result_rows = []
    for key in sorted(set(lookup_a) | set(lookup_b), key=lambda k: (k[0] or "", k[1] or "")):
        norm_vn, fm = key
        ra = lookup_a.get(key)
        rb = lookup_b.get(key)
        a_sev = ra.get("severity") if ra else None
        b_sev = rb.get("severity") if rb else None

        if ra and rb:
            if a_sev is not None and b_sev is None:
                diff, rec = "b_missing", "copy_from_a"
            elif a_sev is not None and b_sev is not None:
                changed = (ra.get("severity") != rb.get("severity") or
                           ra.get("occurrence") != rb.get("occurrence") or
                           ra.get("detection") != rb.get("detection"))
                diff, rec = ("different", "review") if changed else ("same", None)
            else:
                diff, rec = "both_missing", None
        elif ra:
            diff, rec = "a_only", None
        else:
            diff, rec = "b_only", (None if b_sev is not None else "rule_based")

        result_rows.append({
            "norm_key":        norm_vn,
            "failure_mode":    fm,
            "a_variable_name": ra["variable_name"] if ra else None,
            "b_variable_name": rb["variable_name"] if rb else None,
            "a_id":  ra["id"] if ra else None,
            "b_id":  rb["id"] if rb else None,
            "a_severity":   ra.get("severity")   if ra else None,
            "a_occurrence": ra.get("occurrence") if ra else None,
            "a_detection":  ra.get("detection")  if ra else None,
            "a_rpn":        ra.get("rpn")        if ra else None,
            "a_effect":     ra.get("effect_system") if ra else None,
            "a_preventive": ra.get("preventive_action") if ra else None,
            "b_severity":   rb.get("severity")   if rb else None,
            "b_occurrence": rb.get("occurrence") if rb else None,
            "b_detection":  rb.get("detection")  if rb else None,
            "b_rpn":        rb.get("rpn")        if rb else None,
            "diff":         diff,
            "recommendation": rec,
        })

    result_rows.sort(key=lambda r: (_DIFF_ORDER.get(r["diff"], 9), r["norm_key"] or "", r["failure_mode"] or ""))

    summary = {
        "total_a":      len(rows_a),
        "total_b":      len(rows_b),
        "b_missing":    sum(1 for r in result_rows if r["diff"] == "b_missing"),
        "different":    sum(1 for r in result_rows if r["diff"] == "different"),
        "same":         sum(1 for r in result_rows if r["diff"] == "same"),
        "a_only":       sum(1 for r in result_rows if r["diff"] == "a_only"),
        "b_only":       sum(1 for r in result_rows if r["diff"] == "b_only"),
        "both_missing": sum(1 for r in result_rows if r["diff"] == "both_missing"),
    }
    return {"summary": summary, "rows": result_rows}


def apply_compare_patches(patches: list[dict]) -> int:
    """비교에서 선택한 항목을 B 프로젝트에 적용 (A 값으로 덮어씀)"""
    done = 0
    with concurrent.futures.ThreadPoolExecutor(max_workers=15) as ex:
        def apply_one(p):
            patch = {f: p[f"a_{f}"] for f in
                     ("severity", "occurrence", "detection", "effect_system", "preventive_action")
                     if p.get(f"a_{f}") is not None}
            if not patch:
                return False
            return sb_patch("fmea_items", {"id": f"eq.{p['b_id']}"}, patch)
        futs = [ex.submit(apply_one, p) for p in patches]
        for f in concurrent.futures.as_completed(futs):
            if f.result():
                done += 1
    return done


# ── 개념 FMEA (제어기능사양서 기반) ──────────────────────────────────────────
# 개념 FMEA 실패모드 → DB failure_mode 매핑 (DB 제약: CORRUPT/EARLY/LATE/LESS/MORE)
# LESS   = 기능 미수행 (출력이 전혀 없음)
# CORRUPT = 기능 부정확 수행 (잘못된 결과)
# MORE   = 비의도적 수행 (요청 없이 수행)
# LATE   = 지연 수행
CONCEPT_MODES = {
    "LESS":    "{func}이(가) 요구 시점에 수행되지 않음 [기능 미수행]",
    "CORRUPT": "{func}이(가) 부정확하게 수행됨 — 잘못된 결과 생성 [기능 부정확]",
    "MORE":    "{func}이(가) 요청 없이 또는 잘못된 조건에서 수행됨 [비의도적 수행]",
    "LATE":    "{func}이(가) 규정 시간보다 늦게 수행됨 [기능 지연]",
}

_CONCEPT_SEV = {
    "Safety":        {"LESS": 9, "CORRUPT": 8, "MORE": 10, "LATE": 8},
    "Control":       {"LESS": 7, "CORRUPT": 7, "MORE": 8,  "LATE": 6},
    "Monitor":       {"LESS": 5, "CORRUPT": 5, "MORE": 4,  "LATE": 4},
    "Communication": {"LESS": 6, "CORRUPT": 6, "MORE": 6,  "LATE": 7},
    "Default":       {"LESS": 6, "CORRUPT": 6, "MORE": 7,  "LATE": 5},
}
_CONCEPT_OCC = {"LESS": 3, "CORRUPT": 4, "MORE": 2, "LATE": 3}
_CONCEPT_DET = {"LESS": 4, "CORRUPT": 5, "MORE": 7, "LATE": 5}
_CONCEPT_CAUSE = {
    "LESS":    "SW 로직 오류, 입력 신호 누락, 전원/통신 장애",
    "CORRUPT": "연산 로직 결함, 잘못된 파라미터, 센서 오류",
    "MORE":    "상태 머신 오류, 조건 검사 누락, 경쟁 조건(race condition)",
    "LATE":    "CPU 과부하, 통신 지연, 타이머 설정 오류",
}
_CONCEPT_PREVENTIVE = {
    "LESS":    "기능 수행 여부 모니터링, 안전 상태 전환 로직 구현",
    "CORRUPT": "출력 유효성 검사, E2E 보호 적용, 범위 체크",
    "MORE":    "전제조건 검사 강화, 상태 머신 검증, 진입 조건 이중 확인",
    "LATE":    "타임아웃 감시 구현, 태스크 주기 모니터링, alive counter",
}
_CONCEPT_EFFECT = {
    "LESS":    "{func} 미수행으로 인한 시스템 기능 상실",
    "CORRUPT": "{func} 오동작으로 인한 잘못된 제어 출력",
    "MORE":    "{func} 비의도적 수행으로 인한 안전 위험",
    "LATE":    "{func} 지연으로 인한 타임아웃 또는 기능 저하",
}
_CONCEPT_XREF = {
    "LESS":    ["LESS", "MORE", "CORRUPT"],
    "CORRUPT": ["CORRUPT", "MORE"],
    "MORE":    ["MORE", "CORRUPT"],
    "LATE":    ["LATE", "EARLY"],
}


def parse_spec_document(file_data: bytes, filename: str) -> str:
    """제어기능사양서에서 텍스트 추출 (PDF / Excel / Word / TXT)"""
    ext = filename.lower().rsplit(".", 1)[-1]

    if ext == "pdf":
        try:
            import pdfplumber, io
            with pdfplumber.open(io.BytesIO(file_data)) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf.pages)
        except ImportError:
            pass
        try:
            import pypdf, io
            reader = pypdf.PdfReader(io.BytesIO(file_data))
            return "\n".join(pg.extract_text() or "" for pg in reader.pages)
        except ImportError:
            raise RuntimeError("pdfplumber 또는 pypdf 설치 필요")

    if ext in ("xlsx", "xls"):
        try:
            import openpyxl, io
            wb = openpyxl.load_workbook(io.BytesIO(file_data), read_only=True, data_only=True)
            lines = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None and str(c).strip()]
                    if cells:
                        lines.append("\t".join(cells))
            return "\n".join(lines)
        except ImportError:
            raise RuntimeError("openpyxl 설치 필요")

    if ext == "docx":
        try:
            import docx as _docx, io
            doc = _docx.Document(io.BytesIO(file_data))
            lines = [p.text for p in doc.paragraphs if p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        lines.append("\t".join(cells))
            return "\n".join(lines)
        except ImportError:
            raise RuntimeError("python-docx 설치 필요")

    if ext in ("txt", "csv"):
        return file_data.decode("utf-8", errors="replace")

    raise ValueError(f"지원하지 않는 파일 형식: {ext} (PDF/Excel/Word/TXT 지원)")


def extract_functions_ai(text: str, project_name: str) -> list[dict]:
    """Claude AI로 사양서에서 SW 기능 목록 추출 (크레딧 있을 때만)"""
    if not AI_KEY:
        return []
    try:
        import anthropic
        # 타임아웃 15초, 재시도 없음
        client = anthropic.Anthropic(api_key=AI_KEY, timeout=15.0, max_retries=0)
        truncated = text[:8000] if len(text) > 8000 else text
        prompt = (
            f"다음은 '{project_name}' 제어기능사양서 내용입니다.\n\n{truncated}\n\n"
            "이 문서에서 SW 기능(Function) 목록을 추출하세요.\n"
            "각 기능에 대해 JSON 배열로 반환 (마크다운 없이):\n"
            "- name: 기능명 (한국어, 간결하게 10자 이내)\n"
            "- description: 기능 설명 (1-2문장)\n"
            "- category: Safety / Control / Monitor / Communication 중 하나\n\n"
            '[{"name":"기어변속 제어","description":"레버 포지션 기반 변속 명령 생성","category":"Control"},...]'
        )
        resp = client.messages.create(
            model="claude-sonnet-4-6", max_tokens=2048,
            messages=[{"role": "user", "content": prompt}]
        )
        m = re.search(r'\[.*\]', resp.content[0].text.strip(), re.DOTALL)
        if m:
            return json.loads(m.group())
    except Exception as e:
        print(f"AI 기능 추출 오류 (규칙 기반으로 폴백): {e}")
    return []


_BOILERPLATE = (
    "본 문서", "SL CORPORATION", "현대자동차", "기아", "UTC+9", "비밀유지",
    "무단 전재", "정보자산", "Page ", "Export of", "Produced by", "Exported from",
)
_SKIP_SECTIONS = ("I/O Table", "Data Flow", "상세 Description", "개요", "목적",
                  "약어", "참조 문서", "시스템 개요", "외부 인터페이스", "내부 인터페이스")

def extract_functions_heuristic(text: str) -> list[dict]:
    """규칙 기반 기능 추출 — FRS 테이블 + 2단계 섹션 제목 우선"""
    seen: set[str] = set()
    funcs: list[dict] = []

    def add(name: str, desc: str, cat: str):
        name = name.strip().rstrip(":.").strip()
        # 점 리더(TOC) 제거: "기어변속 제어 ......." → "기어변속 제어"
        name = re.sub(r'\s*\.{2,}.*$', '', name).strip()
        # 괄호 이전까지만 사용: "PRA 제어 PRA(Parking..." → "PRA 제어"
        name = re.sub(r'\s+[A-Z가-힣]+\(.*$', '', name).strip()
        # PDF 공백 아티팩트 제거: "ode 동작 모드" → "동작 모드"
        name = re.sub(r'^[a-z]{1,5}\s+([가-힣])', r'\1', name)
        # 숫자/영문 ID만 남는 항목 제외 (예: "ode-24", "FRS-42")
        if re.match(r'^[A-Za-z\-]+[\d\-]+$', name):
            return
        # 문장형(조사 포함) 필터: 기능명이 아닌 문장 제외
        if re.search(r'[의이가을를은는](?:\s|$)', name) and len(name) > 15:
            return
        if not name or len(name) < 3 or name in seen:
            return
        if any(b in name for b in _BOILERPLATE):
            return
        if any(s in name for s in _SKIP_SECTIONS):
            return
        seen.add(name)
        funcs.append({"name": name[:40], "description": desc, "category": cat})

    def guess_cat(n: str) -> str:
        n = n.lower()
        if any(k in n for k in ("안전", "잠금", "방지", "보호", "pra", "파킹")):
            return "Safety"
        if any(k in n for k in ("진단", "고장", "모니터", "감시", "검사", "dtc")):
            return "Monitor"
        if any(k in n for k in ("통신", "can", "송신", "수신", "ota", "네트워크")):
            return "Communication"
        return "Control"

    # 1. FRS 테이블 패턴 (FRS-xxx  기능명  설명)
    for m in re.finditer(r'(FRS-\w+)\s+([가-힣A-Za-z][^\n]{3,40})\s+([가-힣][^\n]{5,})', text):
        name, desc = m.group(2).strip(), m.group(3).strip()[:80]
        add(name, desc, guess_cat(name))

    # 2. N.N 형태 2단계 섹션 (예: "2.1 동작 모드 제어") — 서브섹션(2.1.2) 제외
    for m in re.finditer(r'(?:^|\n)\s*(\d+\.\d+)\s+([가-힣A-Za-z(][^\n]{3,50})', text):
        name = m.group(2).strip()
        add(name, f"{name} 기능 수행", guess_cat(name))

    # 3. 최상위 번호 없는 항목 (예: "○ 기어변속 제어")
    for m in re.finditer(r'(?:^|\n)\s*[○•·]\s+([가-힣][^\n]{4,30})', text):
        name = m.group(1).strip()
        add(name, f"{name} 기능 수행", guess_cat(name))

    return funcs[:80]


def run_concept_pipeline(job_id, jobs, project_name, vehicle_model,
                         doc_data: bytes, doc_filename: str,
                         extra_functions: list[dict]):
    """제어기능사양서 → 개념 FMEA 자동 생성"""
    job = jobs[job_id]

    def log(msg: str, progress: int | None = None):
        job["logs"].append(msg)
        if progress is not None:
            job["progress"] = progress
        print(f"[{job_id[:8]}] {msg}")

    try:
        functions: list[dict] = list(extra_functions)

        # ── 1. 문서 파싱 + 기능 추출 ────────────────────────────────────────
        if doc_data:
            log(f"문서 파싱: {doc_filename}...", 5)
            try:
                doc_text = parse_spec_document(doc_data, doc_filename)
                log(f"  텍스트: {len(doc_text)}자 추출", 12)

                ai_funcs = extract_functions_ai(doc_text, project_name)
                if ai_funcs:
                    log(f"  AI 추출: {len(ai_funcs)}개 기능", 20)
                    functions.extend(ai_funcs)
                else:
                    h_funcs = extract_functions_heuristic(doc_text)
                    log(f"  규칙 추출: {len(h_funcs)}개 기능", 20)
                    functions.extend(h_funcs)
            except Exception as e:
                log(f"  문서 파싱 오류: {e} — 수동 입력만 사용")

        if not functions:
            raise ValueError("기능 목록이 비어 있습니다. 문서를 확인하거나 기능을 직접 입력해주세요.")

        log(f"총 {len(functions)}개 기능 확인", 25)

        # ── 2. 프로젝트 생성 ─────────────────────────────────────────────────
        existing = sb_get("projects", {"name": f"eq.{project_name}", "select": "id"})
        if existing:
            project_id = existing[0]["id"]
            requests.delete(f"{SB_URL}/rest/v1/fmea_items", headers=SB_H,
                            params={"project_id": f"eq.{project_id}"}, verify=False)
            log(f"  기존 프로젝트 초기화: {project_id[:8]}...")
        else:
            r = sb_post("projects", {"name": project_name, "vehicle_model": vehicle_model,
                                      "description": f"개념 FMEA - {project_name}"})
            project_id = (r[0] if isinstance(r, list) else r)["id"]
            log(f"  새 프로젝트: {project_id[:8]}...", 30)

        # SW Units — 카테고리별
        unit_map: dict[str, str] = {}
        for cat in {f.get("category", "Control") for f in functions}:
            r = sb_post("sw_units", {"project_id": project_id, "name": cat})
            unit_map[cat] = (r[0] if isinstance(r, list) else r)["id"]

        # ── 3. FMEA 항목 생성 ────────────────────────────────────────────────
        log("FMEA 항목 생성...", 35)
        db_rows = []
        item_no = 1
        for func in functions:
            fname = func.get("name", "")
            cat   = func.get("category", "Control")
            uid   = unit_map.get(cat)
            sev_map = _CONCEPT_SEV.get(cat, _CONCEPT_SEV["Default"])
            # DB category 제약: Internal / External 만 허용
            db_cat = "External" if cat == "Communication" else "Internal"
            for mode, tmpl in CONCEPT_MODES.items():
                sev = sev_map[mode]
                occ = _CONCEPT_OCC[mode]
                det = _CONCEPT_DET[mode]
                db_rows.append({
                    "project_id":       project_id,
                    "sw_unit_id":       uid,
                    "item_no":          str(item_no),
                    "category":         db_cat,
                    "variable_name":    fname,
                    "variable_type":    cat,   # 기능 카테고리 (Safety/Control/Monitor/Communication)
                    "failure_mode":     mode,
                    "failure_detail":   tmpl.format(func=fname),
                    "potential_cause":  _CONCEPT_CAUSE[mode],
                    "effect_module":    None,
                    "severity":         sev,
                    "occurrence":       occ,
                    "detection":        det,
                    "effect_system":    _CONCEPT_EFFECT[mode].format(func=fname),
                    "preventive_action": _CONCEPT_PREVENTIVE[mode],
                    "status":           "draft",
                    "ai_generated":     False,
                })
                item_no += 1

        log(f"  {len(db_rows)}개 항목 삽입 중...", 50)
        inserted = 0
        for i in range(0, len(db_rows), 200):
            chunk = db_rows[i:i+200]
            r = requests.post(f"{SB_URL}/rest/v1/fmea_items", headers=SB_H,
                              json=chunk, verify=False)
            if r.status_code in (200, 201):
                inserted += len(chunk)
            log(f"  삽입: {min(i+200, len(db_rows))}/{len(db_rows)}")

        # ── 4. 기존 FMEA 교차참조로 S/O/D 보강 ─────────────────────────────
        log("기존 FMEA 교차참조 보강...", 70)
        all_src: list[dict] = []
        for pid in SRC_PROJECT_IDS:
            rows = sb_get("fmea_items", {
                "project_id": f"eq.{pid}", "severity": "not.is.null",
                "select": "variable_name,failure_mode,severity,occurrence,detection,rpn,effect_system,preventive_action",
            })
            all_src.extend(rows)
        norm_lookup = build_sod_lookup(all_src)

        concept_items = sb_get("fmea_items", {
            "project_id": f"eq.{project_id}",
            "select": "id,variable_name,failure_mode",
        })
        patch_batch = []
        for item in concept_items:
            norm = normalize_varname(item["variable_name"])
            for dm in _CONCEPT_XREF.get(item["failure_mode"], ["CORRUPT"]):
                if norm in norm_lookup and dm in norm_lookup[norm]:
                    src = norm_lookup[norm][dm]
                    if (src.get("rpn") or 0) > 0:
                        patch_batch.append((item["id"], {
                            "severity":          src["severity"],
                            "occurrence":        src["occurrence"],
                            "detection":         src["detection"],
                            "effect_system":     src.get("effect_system"),
                            "preventive_action": src.get("preventive_action"),
                        }))
                    break

        xref_done = 0
        if patch_batch:
            with concurrent.futures.ThreadPoolExecutor(max_workers=15) as ex:
                futs = {ex.submit(sb_patch, "fmea_items", {"id": f"eq.{iid}"}, patch): iid
                        for iid, patch in patch_batch}
                for f in concurrent.futures.as_completed(futs):
                    if f.result():
                        xref_done += 1
        log(f"  교차참조 보강: {xref_done}/{len(concept_items)}개", 90)

        log(f"완료! {len(functions)}개 기능 x 4모드 = {inserted}개 항목 생성", 99)
        job["project_id"] = project_id
        job["status"]     = "done"
        job["progress"]   = 100

    except Exception as e:
        import traceback
        job["status"] = "error"
        job["error"]  = str(e)
        log(f"오류: {e}")
        traceback.print_exc()
